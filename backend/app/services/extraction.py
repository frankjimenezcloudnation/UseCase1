"""Layout-light text extraction from the heterogeneous fund/benchmark corpus.

PDFs (FPR, ABTN, Transitieplan, ...) -> pypdf, page-numbered.
DOCX (AllVida spec, Qwik analysis)  -> python-docx.
XLSX (OntologySnapshot, PDC)        -> openpyxl, condensed to concept/property names.
"""

from __future__ import annotations

from pathlib import Path

from pypdf import PdfReader


def extract_pdf(path: Path, max_chars: int) -> str:
    """Return page-tagged text so the model can cite page numbers in ProvisionSource."""
    reader = PdfReader(str(path))
    out: list[str] = []
    total = 0
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        block = f"\n--- page {i} ---\n{text}"
        out.append(block)
        total += len(block)
        if total >= max_chars:
            out.append(f"\n[... truncated at page {i} / {len(reader.pages)} for length ...]")
            break
    return "".join(out).strip()


def extract_docx(path: Path, max_chars: int) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    total = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        parts.append(text)
        total += len(text)
        if total >= max_chars:
            parts.append("[... truncated for length ...]")
            break
    # Tables often hold the actual product specs in these documents.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                line = " | ".join(cells)
                parts.append(line)
                total += len(line)
        if total >= max_chars:
            break
    return "\n".join(parts).strip()


def extract_xlsx(path: Path, max_chars: int) -> str:
    """Condense a workbook to sheet -> the 'Name' column values (concept/property labels)."""
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    total = 0
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c) if c is not None else "" for c in rows[0]]
        name_idx = header.index("Name") if "Name" in header else None
        names: list[str] = []
        for r in rows[1:]:
            if name_idx is not None and name_idx < len(r) and r[name_idx]:
                names.append(str(r[name_idx]))
        if name_idx is None:
            # Fall back to a compact dump of the first few rows.
            preview = ["; ".join(str(c) for c in row if c is not None) for row in rows[1:6]]
            block = f"\n[{ws.title}]\n" + "\n".join(p for p in preview if p)
        else:
            block = f"\n[{ws.title}] ({len(names)} entries): " + ", ".join(names[:200])
        parts.append(block)
        total += len(block)
        if total >= max_chars:
            break
    wb.close()
    return "".join(parts).strip()


def extract_text(path: Path, max_chars: int) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path, max_chars)
    if suffix == ".docx":
        return extract_docx(path, max_chars)
    if suffix in (".xlsx", ".xls"):
        return extract_xlsx(path, max_chars)
    return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
